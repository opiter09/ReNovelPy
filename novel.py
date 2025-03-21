import os
import re
import shutil
import string
import subprocess
import sys
import FreeSimpleGUI as psg
import docx # note: the module is named "python-docx" in pip
import unrpa

weirdData = [
    (b"\xef\xbb\xbf").decode("UTF-8", errors = "ignore")
]

def titleCase(string):
    temp = string.replace("_", " ")
    if (" " not in string):
        temp = temp.replace("-", " ")
    temp2 = ""
    for s in temp.split(" "):
        if (len(s) > 1):
            temp2 = temp2 + s[0].upper() + s[1:].lower()
        elif (len(s) == 1):
            temp2 = temp2 + s[0].upper()
    return(temp2)

def inQuote(string, loc):
    if ((string == "") or (loc >= len(string))):
        return(False)
        
    oneQ = False
    twoQ = False
    for i in range(len(string)):
        if (i == loc):
            if ((string[i] == '"') and (twoQ == True)):
                if ((i == 0) or (string[i - 1] != "\\")):
                    oneQ = False
                    twoQ = False
            elif ((string[i] == "'") and (oneQ == True)):
                if ((i == 0) or (string[i - 1] != "\\")):
                    oneQ = False
                    twoQ = False
            break
        elif ((string[i] == "'") and (twoQ == False)):
            if ((i == 0) or (string[i - 1] != "\\")):
                oneQ = not oneQ
        elif ((string[i] == '"') and (oneQ == False)):
            if ((i == 0) or (string[i - 1] != "\\")):
                twoQ = not twoQ
    return(oneQ or twoQ)
 
def grabQuote(string, loc):
    if ((string == "") or (loc >= len(string)) or (inQuote(string, loc) == False)):
        return("")
    
    res = ""
    for i in range(len(string)):
        test = inQuote(string, i)
        if (test == True):
            res = res + string[i]
        else:
            if (i <= loc):
                res = ""
            else:
                break
    return(res)
    
def firstQuote(string):
    here = 1000 * 1000
    for i in range(len(string)):
        if (inQuote(string, i) == True):
            here = i
            break
    return(grabQuote(string, here))
    
def unquotedCharLoc(string, char):
    check = -1
    for i in range(len(string)):
        if ((string[i] == char) and (inQuote(string, i) == False)):
            check = i
            break
    return(check)
    
path = psg.popup_get_file("Game Executable:", font = "-size 12")
fo = path.split("/")[-2]
fi = path.split("/")[-1]
folder = "./game_" + fo.replace(" ", "_") + "/"
# print(folder)
authorName = psg.popup_get_text("Author Name:", font = "-size 12")
if (authorName == None):
    authorName = ""

if (os.path.exists(folder) == False):
    shutil.copytree(path[0:-(len(fi))] + "game", folder)
    for root, dirs, files in os.walk(folder):
        for file in files:
            if (file.endswith(".rpa") == True):
                rpa = unrpa.UnRPA(filename = os.path.join(root, file), path = folder)
                rpa.extract_files()
                os.remove(os.path.join(root, file))

langs = ["original"]
langRes = "original"
for root, dirs, files in os.walk(folder + "tl/"):
    for dirz in dirs:
        if (dirz.lower() != "none") and (os.path.exists(folder + "tl/" + dirz) == True):
            langs.append(dirz)
if (len(langs) > 1):
    layout = []
    for i in range(len(langs)):
        layout = layout + [[psg.Button(titleCase(langs[i].strip()), key = "choice_" + langs[i])]]
    window = psg.Window("", layout, grab_anywhere = True, resizable = True, font = "-size 12")
    while True:
        event, values = window.read()
        # See if user wants to quit or window was closed
        if (event == psg.WINDOW_CLOSED) or (event == "Quit"):
            break
        elif (event.startswith("choice_") == True):
            langRes = event.split("_")[1]
            break
    window.close()

nameSwaps = {}
if (langRes != "original"):
    if (os.path.exists(folder + "tl/" + langRes + "/names.rpy") == True):
        f = open(folder + "tl/" + langRes + "/names.rpy", "rb")
        r = f.read().decode("UTF-8", errors = "ignore")
        f.close()
        for w in weirdData:
            r = r.replace(w, "")
        lines = r.split("\n")
        for i in range(len(lines) - 1):
            temp = lines[i].strip()
            if (temp.startswith("old") == True):
                nameSwaps[firstQuote(temp)] = firstQuote(temp)
                for j in range(i + 1, len(lines)):
                    temp2 = lines[j].strip()
                    if ((temp2.startswith("new") == True) and (firstQuote(temp2) != "")):
                        nameSwaps[firstQuote(temp)] = firstQuote(temp2)
                        break
 
layout = [
    [ psg.Button("No Images", key = "none") ],
    [ psg.Button("Characters Only", key = "chara") ],
    [ psg.Button("Chars. and Selected", key = "select") ],
    [ psg.Button("All Images", key = "all") ]
]

window = psg.Window("", layout, grab_anywhere = True, resizable = True, font = "-size 12")

imageChoice = "none"
while True:
    event, values = window.read()
    # See if user wants to quit or window was closed
    if (event == psg.WINDOW_CLOSED) or (event == "Quit"):
        break
    elif (event in ["none", "chara", "select", "all"]):
        imageChoice = event
        break
window.close()

selectFiles = []
if (imageChoice == "select"):
    check = 0
    while (check == 0):
        path2 = psg.popup_get_file("Included Image (Leave Blank to Finish):", font = "-size 12")
        if (path2 == None):
            path2 = ""
        if (path2 == ""):
            check = 1
        else:
            san = path2.replace("\\", "/")
            if ("/game/" in san):
                san = os.getcwd().replace("\\", "/") + "/" + folder[2:] + san.split("/game/")[1]
            if ("/tl/" in san):
                langLen = len(san.split("/tl/")[1].split("/")[0])
                san = san.split("/tl/")[0] + san.split("/tl/")[1][langLen:]
            selectFiles.append(san)
# print(selectFiles)

labels = ["start"]
f = open(folder + "screens.rpy", "rb")
r = f.read().decode("UTF-8", errors = "ignore") # python is weird about Japanese etc. when using "rt"
f.close()
for w in weirdData:
    r = r.replace(w, "")
for l in r.split("\n"):
    if ("action Start(" in l):
        labels.append(firstQuote(l.split("action Start(")[1]))
# print(labels)

title = ""
f = open(folder + "options.rpy", "rb")
r = f.read().decode("UTF-8", errors = "ignore")
f.close()
for w in weirdData:
    r = r.replace(w, "")
for l in r.split("\n"):
    if (("config.name" in l) or ("config.window_title" in l)):
        title = firstQuote(l)
        break
titleU = title.strip().upper()
# print(title)
if (langRes != "original"):
    titleU = titleU + " [" + langRes.strip().upper() + "]"

combined = "" # to avoid looping through files all the time
for root, dirs, files in os.walk(folder):
    for file in files:
        if (file.endswith(".rpy") == True):
            f = open(os.path.join(root, file), "rb")
            r = f.read()
            f.close()
            # if (b"label start:" in r):
                # print(file)
            combined = combined + "\n\n" + r.decode("UTF-8", errors = "ignore")
for w in weirdData:
    combined = combined.replace(w, "")
# print(combined[0:50])

combLSpaced = list(combined.split("\n")).copy()
combL = [x.strip() for x in combLSpaced]
for i in range(len(combL)):
    l = combL[i]
    here = unquotedCharLoc(l, "#")
    if (here >= 0):
        combL[i] = combL[i][0:here].strip()

if (langRes != "original"):
    for root, dirs, files in os.walk(folder + "tl/" + langRes):
        for file in files:
            if ((file.endswith(".rpy") == True) and (file not in ["common.rpy", "names.rpy", "screens.rpy", "text history.rpy"])):
                f = open(os.path.join(root, file), "rb")
                r = f.read().decode("UTF-8", errors = "ignore")
                f.close()
                for w in weirdData:
                    r = r.replace(w, "")
                lines = r.split("\n")
                oldStrings = []
                newStrings = []
                for i in range(len(lines)):
                    temp = lines[i].strip()
                    if ((temp.startswith("translate") == True)):
                        # print(oldStrings)
                        # print(newStrings)
                        for j in range(min(len(oldStrings), len(newStrings))):
                            first = oldStrings[j]
                            here = unquotedCharLoc(first, "#")
                            if (here >= 0):
                                first = first[0:here].strip()
                            new = newStrings[j]
                            here = unquotedCharLoc(new, "#")
                            if (here >= 0):
                                new = new[0:here].strip()
                            for k in range(len(combL)):
                                if (combL[k] == first):
                                    combL[k] = new
                        if ("strings" not in temp):
                            oldStrings = [""]
                            newStrings = [""]
                        else:
                            oldStrings = []
                            newStrings = []
                    elif (len(oldStrings) > 0):
                        if ((temp != "") and (temp[0] == "#") and (('"' in temp) or ("'" in temp))):
                            oldStrings.append(temp[1:].strip())
                        elif (('"' in temp) or ("'" in temp)):
                            newStrings.append(temp.strip())
                if (len(oldStrings) > 0):
                    for j in range(min(len(oldStrings), len(newStrings))):
                        first = oldStrings[j]
                        here = unquotedCharLoc(first, "#")
                        if (here >= 0):
                            first = first[0:here].strip()
                        new = newStrings[j]
                        here = unquotedCharLoc(new, "#")
                        if (here >= 0):
                            new = new[0:here].strip()
                        for k in range(len(combL)):
                            if (combL[k] == first):
                                combL[k] = new
    for root, dirs, files in os.walk(folder + "tl/" + langRes): # separated to avoid substring overwriting
        for file in files:
            if ((file.endswith(".rpy") == True) and (file not in ["common.rpy", "names.rpy", "screens.rpy", "text history.rpy"])):
                f = open(os.path.join(root, file), "rb")
                r = f.read().decode("UTF-8", errors = "ignore")
                f.close()
                for w in weirdData:
                    r = r.replace(w, "")
                lines = r.split("\n")
                for i in range(len(lines)):
                    temp = lines[i].strip()
                    if ((i != (len(lines) - 1)) and (len(temp) >= 4) and (temp[0:4] in ["old ", 'old"', "old'"])):
                        for j in range(i + 1, len(lines)):
                            temp2 = lines[j].strip()
                            if ((len(temp2) >= 4) and (temp2[0:4] in ["new ", 'new"', "new'"])):
                                first = firstQuote(temp)
                                new = firstQuote(temp2)
                                for k in range(len(combL)):
                                    if ((first in combL[k]) and (combL[k].startswith("label") == False)):
                                        combL[k] = combL[k].replace(first, new)
                                break        

def findSprite(image):
    global combL
    global folder

    if (image == ""):
        return("")

    sprite = ""
    for i in range(len(combL)):
        l2 = combL[i]
        if ((l2.startswith("image ") == True) and (((" " + image + " ") in l2) or ((" " + image + "=") in l2) or ((" " + image + ":") in l2))):
            if ("Movie(" in l2):
                break
            elif ((".png" in l2) or (".jpg" in l2)):
                j = i
            else:
                j = i
                while ((j < (len(combL) - 1)) and (".png" not in combL[j]) and (".jpg" not in combL[j])) or (combL[j][0] == "#"):
                    j = j + 1
            for k in range(len(combL[j])):
                small = grabQuote(combL[j], k)
                if ((small.endswith(".png") == True) or (small.endswith(".jpg") == True)):
                    if (os.path.exists(folder + small) == True):
                        sprite = folder + small
                    else:
                        if (os.path.exists(folder + "images/" + small) == True):
                            sprite = folder + "images/" + small
                        else:
                            sprite = ""
                    # print(sprite)
                    break
            break
    return(sprite)
    
new = docx.Document()
head = new.add_heading(titleU, 0)
head.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
if (authorName != ""):
    head = new.add_heading("BY " + authorName.upper(), 0)
    head.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
new.add_page_break()

section = new.sections[0]
pageWidth = section.page_width - section.left_margin - section.right_margin

nameVars = {}
inputVars = {}
usedNames = {}
for i in range(len(combL)):
    l = combL[i]
    var = ""
    res = ""
    if ("renpy.input(" in l):
        var = l.split("=")[0].strip()
        res = titleCase(var)
        if (i < (len(combL) - 1)):
            for j in range(i + 1, len(combL)):
                if ((combL[j] != "") and (firstQuote(combL[j]) != "")):
                    if (unquotedCharLoc(combL[j], "=") >= 0):
                        var = combL[j].split("=")[0].strip()
                        res = firstQuote(combL[j])
                        break
                    else:
                        break
        res = res.replace("\\[", "[").replace("\\]", "]")
        res = res.replace("\\{", "<").replace("\\}", ">")
        for j in range(len(res)):
            if ((j < (len(res) - 1)) and (res[j] == "}")):
                if (res[j + 1] != "{"):
                    res = res[(j + 1):].split("{")[0]
                    break
        nameVars[var] = res
        inputVars[var] = res
        usedNames[res] = ""
for i in range(len(combL)):
    l = combL[i]
    var = ""
    name = ""
    if ("Character(" in l):
        afterString = l.split("Character(")[1].strip()
        if ((afterString != "") and (afterString[0] in ['"', "'"])):
            name = firstQuote(afterString)
        elif ("name=" in afterString.replace("name =", "name=")):
            name = firstQuote(afterString.replace("name =", "name=").split("name=")[1])
        elif (afterString.startswith("None") == True):
            name = "~|NONE|~" # surely no-one actually uses this, right?
        elif (((afterString == "") or (afterString[-1] == ",")) and (i < (len(combL) - 1))):
            for j in range(i + 1, len(combL)):
                if (combL[j] != ""):
                    if ((combL[j][0] in ['"', "'"]) or (combL[j].replace("name =", "name=").startswith("name=") == True)):
                        name = firstQuote(combL[j])
                        if ("name=None" in combL[j].replace(" ", "")):
                            name = "~|NONE|~"
                        break
                    elif (combL[j].startswith("None") == True):
                        name = "~|NONE|~"
                        break
                    elif (combL[j][-1] == ")"):
                        break
        if (name in nameSwaps.keys()):
            name = nameSwaps[name]
        for iv in inputVars.keys():
            name = name.replace("[" + iv + "]", inputVars[iv])
        name = name.replace("\\[", "[").replace("\\]", "]")
        name = name.replace("\\{", "<").replace("\\}", ">")
        for j in range(len(name)):
            if ((j < (len(name) - 1)) and (name[j] == "}")):
                if (name[j + 1] != "{"):
                    name = name[(j + 1):].split("{")[0]
                    break
        var = l[7:].split("=")[0].strip()
        nameVars[var] = name
        image = ""
        j = i
        while (j < (len(combL) - 1)) and ((combL[j] == "") or (combL[j][-1] != ")") or (j == i)) and ("ImageReference(" not in combL[j]) and ("image=" not in combL[j].replace("image =", "image=")):
            j = j + 1
        if ("ImageReference(" in combL[j]):
            image = firstQuote(combL[j].split("ImageReference(")[1])
        elif ("image=" in combL[j].replace("image =", "image=")):
            image = firstQuote(combL[j].replace("image =", "image=").split("image=")[1])
        sprite = ""
        if (image != ""):
            # print(image)
            sprite = findSprite(image)
        if (((name not in usedNames.keys()) or (usedNames[name] == "")) and (name.replace("?", "") != "")):
            if ((sprite not in usedNames.values()) or (sprite == "")):
                usedNames[name] = sprite
    elif ("nvl_narrator" in l):
        var = l[7:].split("=")[0].strip()
        nameVars[var] = ""
# print(usedNames)

if (imageChoice != "none"):
    new.add_heading("Dramatis Personæ", 0)
    nk = list(usedNames.keys()).copy()
    nk.sort()
    for n in nk:
        if (n != "~|NONE|~"):
            new.add_heading(n, 1)
            if (usedNames[n] != ""):
                sp = usedNames[n]
                if (os.path.exists(sp.replace(folder, folder + "tl/" + langRes + "/")) == True):
                    sp = sp.replace(folder, folder + "tl/" + langRes + "/")
                new.add_picture(sp, width = int(pageWidth * 0.5))
    new.add_page_break()

allImageVars = []
otherImageVars = []
for l in combL:
    if (l.startswith("image ") == True):
        temp = l[6:].split(":")[0].split("=")[0].strip()
        allImageVars.append(temp)
        check = 0
        numberList = ["0", "1", "2", "3", "4", "5", "6", "7", "8", "9"]
        for n in (list(nameVars.keys()).copy() + list(nameVars.values()).copy()):
            endNumLost = re.sub(r"[0-9]*", "", temp[::-1], count = 1)[::-1]
            if (temp.lower().replace("_", " ").replace("-", " ") == n.lower()):
                check = 1
            elif (temp.lower().replace("_", " ").replace("-", " ").startswith(n.lower() + " ") == True):
                check = 1
            elif (temp.lower().replace("_", " ").replace("-", " ").endswith(" " + n.lower()) == True):
                check = 1
            elif ((temp.lower().startswith(n.lower()) == True) and (len(temp) > len(n))):
                if ((temp[len(n)] in numberList) or ((temp[len(n) - 1] in string.ascii_lowercase) and temp[len(n)] in string.ascii_uppercase)):
                    check = 1
            elif ((endNumLost.lower().endswith(n.lower()) == True) and (len(endNumLost) > len(n))):
                if (((endNumLost[-(len(n) + 1)] == " ") and (temp[-1] in numberList)) or ((endNumLost[-len(n)] in string.ascii_uppercase) and (endNumLost[-(len(n) + 1)] in string.ascii_lowercase))):
                    check = 1
        if (check == 0):
            otherImageVars.append(temp)
allImageVars.sort() # shorter names come first
otherImageVars.sort()
# print(allImageVars)

p = ""
def handleTags(string, boldName):
    global p
    
    if (string == ""):
        return

    drop = string
    if (boldName == True):
        if ((drop[0] == "“") and (drop[-1] == "”")):
            drop = drop[1:-1]
        elif ((len(drop) >= 2) and (drop[0] in ['"', "'"]) and (drop[0] == drop[-1])):
            drop = drop[1:-1]
        elif ((len(drop) >= 4) and (drop[0:2] in ['\\"', "\\'"]) and (drop[0:2] == drop[-2:])):
            drop = drop[2:-2]
    drop = drop.replace("\\n", " ")
    for iv in inputVars.keys():
        drop = drop.replace("[" + iv + "]", inputVars[iv])
    drop = drop.replace("\\[", "[").replace("\\]", "]")
    drop = drop.replace("\\{", "<").replace("\\}", ">")
    drop = drop.replace('\\"','"').replace("\\'", "'")
        
    tags =  re.split(r"([\{\}])", drop) # thank you Stack Exchange
    runs = []
    for t in tags:
        runs.append([t])
    lastSize = ""
    for i in range(len(tags) - 2):
        if (tags[i] == "{"):
            runs[i].append("g" * 1000)
            runs[i + 1].append("g" * 1000)
            runs[i + 2].append("g" * 1000)
            if (tags[i + 1][0] != "/"):
                safe = tags[i + 1].replace(" ", "")
                # b is bold, i is italics, s is strikethrough, u is underline, plain removes all 4, space is x spaces, size does math
                if (safe.split("=")[0] in ["b", "i", "s", "u", "plain", "size"]):
                    for j in range(i + 1, len(tags)):
                        runs[j].append(safe)
                    if (safe.split("=")[0] == "size"):
                        lastSize = safe
                elif (safe.split("=")[0] == "space"):
                    runs[i + 1][0] = " " # * int(safe.split("=")[1])
            else:
                safe = tags[i + 1].replace(" ", "")
                for j in range(i + 1, len(tags)):
                    try:
                        runs[j].remove(safe[1:])
                    except ValueError:
                        if (safe == "/size"):
                            runs[j].remove(lastSize)
                            
    for r in runs:
        if ((r[0] != "") and ((("g" * 1000) not in r) or (r[0].replace(" ", "") == ""))):
            thing = p.add_run(r[0])
            thing.font.size = docx.shared.Pt(12)
            if ("b" in r):
                thing.bold = True
            if ("i" in r):
                thing.italic = True
            if ("s" in r):
                thing.font.strike = True
            if ("u" in r):
                thing.underline = True
            if ("plain" in r):
                thing.bold = False
                thing.italic = False
                thing.font.strike = False
                thing.underline = False
            for val in r:
                if (val.startswith("size") == True):
                    if ("*" in val):
                        thing.font.size = max(8, int(thing.font.size * float(val.split("*")[1])))
                    elif ("+" in val):
                        thing.font.size = thing.font.size + docx.shared.Pt(int(val.split("+")[1]))
                    elif ("-" in val):
                        thing.font.size = max(8, thing.font.size - docx.shared.Pt(int(val.split("-")[1])))
                    else:
                        thing.font.size = docx.shared.Pt(int(val.split("=")[1]))      

curr = ""
menuChoices = []
skipInd = -1
returnInd = -1
returnCurr = ""
returnMC = []
returnSI = -1
first = 0
triple = False
for lab in labels:
    curr = lab
    while ("\n" + "label " + curr + ":") in combined:
        # print(curr)
        if (first == 0):
            first = 1
        else:
            new.add_page_break()
        new.add_heading(titleCase(curr), 0)
        if (returnInd != -2):
            ind = 0
            for i in range(len(combL)):
                if (combL[i].startswith("label " + curr + ":") == True):
                    ind = i + 1
                    break
        else:
            returnInd = -1
        while True:
            line = combL[ind]
            if ((line == "") or (line.startswith("#") == True)):
                ind = ind + 1
                if (ind == len(combL)):
                    curr = "g" * 1000
                    # print("finalBlank")
                    break
                else:
                    continue
            # the below has a billion weird conditions to ensure that various hypothetical permutations of """ placement still work
            if (('"""' in line) or ("'''" in line)):
                triple = not triple
                if ((triple == True) and (line.replace("'''", '"""').split('"""')[1] == "")):
                    ind = ind + 1
                    if (ind == len(combL)):
                        curr = "g" * 1000
                        break
                    else:
                        continue
                elif ((triple == False) and (line.replace("'''", '"""').split('"""')[0] == "")):
                    ind = ind + 1
                    if (ind == len(combL)):
                        curr = "g" * 1000
                        break
                    else:
                        continue
            if ((triple == True) or ((('"""' in line) or ("'''" in line)) and (line.replace("'''", '"""').split('"""')[0] != ""))):
                p = new.add_paragraph()
                if (('"""' in line) or ("'''" in line)):
                    if (triple == True):
                        handleTags(line.replace("'''", '"""').split('"""')[1], False)
                    else:
                        handleTags(line.replace("'''", '"""').split('"""')[0], False)
                else:
                    handleTags(line, False)
                if (line.replace("'''", '"""').endswith('"""') == True):
                    triple = False
                ind = ind + 1
                if (ind == len(combL)):
                    curr = "g" * 1000
                    break
                else:
                    continue
            if ((line.startswith("play sound ") == True) or (line.startswith("play audio ") == True)):
                soundList = []
                qCheck = False
                whole = ""
                if ((unquotedCharLoc(line, "[") >= 0) and (unquotedCharLoc(line, "]") == -1)):
                    for i in range(ind + 1, len(combL)):
                        whole = whole + combL[i] + "\n"
                        if (unquotedCharLoc(combL[i], "]") >= 0):
                            ind = i
                            break
                else:
                    whole = line
                for i in range(len(whole)):
                    if (inQuote(whole, i) == True):
                        if (qCheck == False):
                            qCheck = True
                            soundList.append(grabQuote(whole, i))
                    else:
                        qCheck = False
                for s in soundList:
                    # if (len(soundList) > 1):
                        # print(s)
                    sound = s
                    if (sound[0] == "<"):
                        angleLen = len(sound.split(">")[0])
                        sound = sound[(angleLen + 1):]
                    if ((len(sound) >= 4) and (sound[-4] == ".")):
                        sound = sound[0:-4]
                    sound = sound.replace("\\", "/").split("/")[-1]
                    sound = sound.strip()
                    if (sound != ""):
                        p = new.add_paragraph()
                        r = p.add_run("Sound: " + titleCase(sound))
                        r.font.size = docx.shared.Pt(14)
                        r.italic = True
            elif (line.startswith("scene ") == True):
                sprite = ""
                temp = ""
                for v in allImageVars:
                    if ((" " + v) in line[5:]):
                        temp = v # don't break so longer names trump shorter ones
                sprite = findSprite(temp)
                if (sprite == ""):
                    chop = line[6:]
                    for func in ["with", "as", "at", "behind", "onlayer", "zorder"]:
                        chop = chop.split(" " + func + " ")[0]
                    if (chop[-1] == ":"):
                        chop = chop[0:-1]
                    for root, dirs, files in os.walk(folder):
                        for file in files:
                            if ((len(file) >= 4) and (file[0:-4] == chop) and (file[-4:] in [".png", ".jpg"])):
                                sprite = os.path.join(root, file)
                                break
                if (sprite != ""):
                    # print(sprite)
                    if ((imageChoice == "all") or ((imageChoice == "select") and ((os.getcwd() + "/" + sprite[2:]).replace("\\", "/") in selectFiles))):
                        if (os.path.exists(sprite.replace(folder, folder + "tl/" + langRes + "/")) == True):
                            sprite = sprite.replace(folder, folder + "tl/" + langRes + "/")
                        new.add_picture(sprite, width = pageWidth)
                    else:
                        p = new.add_paragraph()
                        r = p.add_run("Scene: " + titleCase(sprite.replace("\\", "/").split("/")[-1][0:-4]))
                        r.font.size = docx.shared.Pt(14)
                        r.italic = True
            elif ((line.startswith("show ") == True) and (line.startswith("show text ") == False) and (line.startswith("show screen ") == False)):
                sprite = ""
                temp = ""
                bad = False
                for v in allImageVars:
                    if ((" " + v) in line[4:]):
                        temp = v # don't break so longer names trump shorter ones
                if (temp in otherImageVars):
                    sprite = findSprite(temp)
                elif (temp != ""):
                    bad = True
                if ((sprite == "") and (bad == False)):
                    chop = line[5:]
                    for func in ["with", "as", "at", "behind", "onlayer", "zorder"]:
                        chop = chop.split(" " + func + " ")[0]
                    if (chop[-1] == ":"):
                        chop = chop[0:-1]
                    for root, dirs, files in os.walk(folder):
                        for file in files:
                            if ((len(file) >= 4) and (file[0:-4] == chop) and (file[-4:] in [".png", ".jpg"])):
                                sprite = os.path.join(root, file)
                                check = 0
                                endNumLost = re.sub(r"[0-9]*", "", file[0:-4][::-1], count = 1)[::-1]
                                for n in (list(nameVars.keys()).copy() + list(nameVars.values()).copy()):
                                    if (file.lower().replace("_", " ").replace("-", " ") == n.lower()):
                                        check = 1
                                    elif (file.lower().replace("_", " ").replace("-", " ").startswith(n.lower() + " ") == True):
                                        check = 1
                                    elif (file[0:-4].lower().replace("_", " ").replace("-", " ").endswith(" " + n.lower()) == True):
                                        check = 1
                                    elif ((file.lower().startswith(n.lower()) == True) and (len(file[0:-4]) > len(n))):
                                        if ((file[len(n)] in numberList) or ((file[len(n) - 1] in string.ascii_lowercase) and file[len(n)] in string.ascii_uppercase)):
                                            check = 1
                                    elif ((endNumLost.lower().endswith(n.lower()) == True) and (len(endNumLost) > len(n))):
                                        if (((endNumLost[-(len(n) + 1)] == " ") and (file[0:-4][-1] in numberList)) or ((endNumLost[-len(n)] in string.ascii_uppercase) and (endNumLost[-(len(n) + 1)] in string.ascii_lowercase))):
                                            # print(sprite)
                                            check = 1
                                if (check == 1):
                                    sprite = ""
                                break
                if (sprite != ""):
                    if ((imageChoice == "all") or ((imageChoice == "select") and ((os.getcwd() + "/" + sprite[2:]).replace("\\", "/") in selectFiles))):
                        if (os.path.exists(sprite.replace(folder, folder + "tl/" + langRes + "/")) == True):
                            sprite = sprite.replace(folder, folder + "tl/" + langRes + "/")
                        new.add_picture(sprite, width = pageWidth)
                    else:
                        p = new.add_paragraph()
                        r = p.add_run("Visual: " + titleCase(sprite.replace("\\", "/").split("/")[-1][0:-4]))
                        r.font.size = docx.shared.Pt(14)
                        r.italic = True
            elif (line.startswith("jump ") == True):
                curr = line[5:]
                # print("jump")
                menuChoices = []
                skipInd = -1
                break
            elif (line.startswith("call ") == True):
                returnCurr = curr
                returnInd = ind + 1
                curr = line[5:].split(" from ")[0]
                # print("call")
                returnMC = menuChoices
                returnSI = skipInd
                if (" from " in line):
                    returnCurr = line.split(" from ")[1]
                    returnInd = -1
                    returnMC = []
                    returnSI = -1
                menuChoices = []
                skipInd = -1
                break
            elif (line.startswith("menu:") == True):
                menuChoices = []
                spacesOld = 0
                for s in combLSpaced[ind]:
                    if (s == " "):
                        spacesOld = spacesOld + 1
                    elif (s == "\t"):
                        spacesOld = spacesOld + 4
                    else:
                        break
                for j in range(ind + 1, len(combL)):
                    spacesNew = 0
                    for s in combLSpaced[j]:
                        if (s == " "):
                            spacesNew = spacesNew + 1
                        elif (s == "\t"):
                            spacesNew = spacesNew + 4
                        else:
                            break
                    if ((combL[j] != "") and (spacesNew <= spacesOld)):
                        skipInd = j
                        break
                    elif ((combL[j] != "") and (combL[j][0] in ['"', "'"]) and (combL[j][-1] == ":")):
                        menuChoices.append([combL[j][1:-2], j + 1])
                layout = []
                for k in range(len(menuChoices)):
                    layout = layout + [[psg.Button(menuChoices[k][0], key = "choice_" + str(k))]]
                window = psg.Window("", layout, grab_anywhere = True, resizable = True, font = "-size 12")
                res = 0
                while True:
                    event, values = window.read()
                    # See if user wants to quit or window was closed
                    if (event == psg.WINDOW_CLOSED) or (event == "Quit"):
                        break
                    elif (event.startswith("choice_") == True):
                        res = int(event.split("_")[1])
                        break
                window.close()
                ind = menuChoices[res][1]
                # print(ind)
                # print(skipInd)
                continue    
            elif (line.startswith("show text ") == True):
                if (firstQuote(line) != ""):
                    p = new.add_paragraph()
                    handleTags(firstQuote(line), False)
            elif (line == "return"):
                if (returnCurr != ""):
                    ind = returnInd
                    curr = returnCurr
                    menuChoices = returnMC
                    skipInd = returnSI
                    if (returnInd != -1):
                        returnInd = -2
                    returnCurr = ""
                    returnMC = []
                    returnSI = -1
                    break
                else:
                    curr = "g" * 1000
                    # print("finalReturn")
                    break
            elif ("renpy.full_restart()" in line):
                curr = "g" * 1000
                # print("finalReturn")
                break
            elif (line.startswith("label ") == True):
                curr = line[6:-1]
                break
            elif (line[0] in ['"', "'"]):
                if (firstQuote(line) != ""):
                    p = new.add_paragraph()
                    handleTags(firstQuote(line), False)
            else:
                theKeys = list(nameVars.keys()).copy()
                theKeys.sort() # shorter names come first
                temp = ""
                check = 0
                for k in theKeys:
                    if ((line.startswith(k + " ") == True) or (line.startswith(k + '"') == True) or (line.startswith(k + "'") == True)):
                        temp = k # don't break so longer names trump shorter ones
                if ((temp != "") and (firstQuote(line) != "")):
                    p = new.add_paragraph()
                    if (nameVars[temp] != "~|NONE|~"):
                        if (nameVars[temp] != ""):
                            check = 1
                            name = p.add_run(nameVars[temp] + ": ")
                            name.bold = True
                            name.font.size = docx.shared.Pt(12)
                    if (check == 0):
                        handleTags(firstQuote(line), False)
                    else:
                        handleTags(firstQuote(line), True)
            ind = ind + 1
            if ((ind == len(combL)) and (line.startswith("jump ") == False)):
                curr = "g" * 1000
                # print("finalNormal")
                break
            elif ((len(menuChoices) > 0) and (ind in [x[1] for x in menuChoices])):
                ind = skipInd
                menuChoices = []
                skipInd = -1
    # print(curr + " out")

new.save("./" + title.replace(" ", "_") + ".docx")